from __future__ import annotations

import json
import math
import os
import subprocess
import sys
import threading
import time
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import logging

import numpy as np

logger = logging.getLogger(__name__)


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DATA_ROOT = PROJECT_ROOT / "data" / "top_ready"
OUTPUT_ROOT = PROJECT_ROOT / "outputs"
RELATION_SCRIPT = PROJECT_ROOT / "relation_evgat" / "run_top_ready_relation_gat.py"


@dataclass
class TrainJob:
    job_id: str
    dataset: str
    status: str = "queued"
    created_at: float = field(default_factory=time.time)
    updated_at: float = field(default_factory=time.time)
    log_path: str | None = None
    output_dir: str | None = None
    error: str | None = None


JOBS: dict[str, TrainJob] = {}
JOBS_LOCK = threading.Lock()


def _read_json(path: Path, fallback: Any = None) -> Any:
    if not path.exists():
        return fallback
    return json.loads(path.read_text(encoding="utf-8"))


def _load_npy(path: Path, fallback: np.ndarray | None = None) -> np.ndarray:
    if path.exists():
        return np.load(path, allow_pickle=False)
    if fallback is not None:
        return fallback
    raise FileNotFoundError(str(path))


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        value = float(value)
        if math.isnan(value) or math.isinf(value):
            return default
        return value
    except Exception:
        return default


def _downsample(items: list[dict[str, Any]], limit: int = 320) -> list[dict[str, Any]]:
    if len(items) <= limit:
        return items
    step = max(1, math.ceil(len(items) / limit))
    return items[::step]


def available_datasets() -> list[dict[str, Any]]:
    datasets = []
    if not DATA_ROOT.exists():
        return datasets
    for dataset_dir in sorted(p for p in DATA_ROOT.iterdir() if p.is_dir()):
        summary = _read_json(dataset_dir / "summary.json", {})
        result_dir = latest_result_dir(dataset_dir.name)
        datasets.append(
            {
                "id": dataset_dir.name,
                "name": dataset_dir.name,
                "num_features": summary.get("num_features"),
                "train_shape": summary.get("train_shape"),
                "test_shape": summary.get("test_shape"),
                "anomaly_ratio": summary.get("test_anomaly_ratio"),
                "has_outputs": result_dir is not None,
                "result_tag": result_dir.name if result_dir else None,
            }
        )
    return datasets


def latest_result_dir(dataset: str) -> Path | None:
    base = OUTPUT_ROOT / "top_ready_relation_gat" / dataset
    preferred = base / "full_joint"
    if preferred.exists():
        return preferred
    if not base.exists():
        return None
    candidates = [p for p in base.iterdir() if p.is_dir() and (p / "summary.json").exists()]
    if not candidates:
        return None
    return max(candidates, key=lambda p: p.stat().st_mtime)


def dataset_paths(dataset: str) -> dict[str, Path]:
    data_dir = DATA_ROOT / dataset
    result_dir = latest_result_dir(dataset)
    if not data_dir.exists():
        raise FileNotFoundError(f"Dataset not found: {dataset}")
    if result_dir is None:
        raise FileNotFoundError(f"No Relation-EVGAT outputs found for dataset: {dataset}")
    return {"data": data_dir, "result": result_dir}


def load_dataset_summary(dataset: str) -> dict[str, Any]:
    data_dir = DATA_ROOT / dataset
    summary = _read_json(data_dir / "summary.json", {})
    columns = _read_json(data_dir / "columns.json", [])
    summary["columns"] = columns
    return summary


def load_result_bundle(dataset: str) -> dict[str, Any]:
    paths = dataset_paths(dataset)
    result_dir = paths["result"]
    data_dir = paths["data"]
    summary = _read_json(result_dir / "summary.json", {})
    columns = _read_json(data_dir / "columns.json", [])
    label = _load_npy(data_dir / "test_label.npy")
    test = _load_npy(data_dir / "test.npy")
    times = _load_npy(result_dir / "times.npy").astype(int)
    node_score = _load_npy(result_dir / "node_score.npy")
    edge_score = _load_npy(result_dir / "edge_score.npy", np.zeros_like(node_score))
    joint_score = _load_npy(result_dir / "joint_score.npy")
    node_errors = _load_npy(result_dir / "node_errors.npy")
    return {
        "data_dir": data_dir,
        "result_dir": result_dir,
        "summary": summary,
        "columns": columns,
        "label": label,
        "test": test,
        "times": times,
        "node_score": node_score,
        "edge_score": edge_score,
        "joint_score": joint_score,
        "node_errors": node_errors,
    }


def threshold_from_summary(summary: dict[str, Any], scores: np.ndarray) -> float:
    joint = summary.get("joint_node_edge", {})
    threshold = joint.get("threshold") or joint.get("best_f1_threshold")
    if threshold is None:
        threshold = float(np.quantile(scores, 0.95))
    return _safe_float(threshold, float(np.quantile(scores, 0.95)))


def event_intervals(label: np.ndarray, times: np.ndarray | None = None) -> list[dict[str, int]]:
    if times is None:
        aligned = label.astype(int)
        index_values = np.arange(len(label))
    else:
        index_values = times.astype(int)
        aligned = label[index_values].astype(int)
    events = []
    start = None
    event_id = 1
    for idx, value in enumerate(aligned):
        if value == 1 and start is None:
            start = idx
        if start is not None and (value == 0 or idx == len(aligned) - 1):
            end = idx - 1 if value == 0 else idx
            events.append(
                {
                    "event_id": event_id,
                    "start": int(index_values[start]),
                    "end": int(index_values[end]),
                    "aligned_start_index": int(start),
                    "aligned_end_index": int(end),
                    "duration": int(end - start + 1),
                }
            )
            event_id += 1
            start = None
    return events


def root_cause_events(dataset: str) -> list[dict[str, Any]]:
    path = OUTPUT_ROOT / "root_cause" / dataset / "event_root_cause_candidates.json"
    rows = _read_json(path, [])
    if isinstance(rows, list):
        return rows
    return []


def selected_event(dataset: str, event_id: int | None = None) -> dict[str, Any]:
    rc_events = root_cause_events(dataset)
    if rc_events:
        if event_id is None:
            event_id = int(rc_events[0].get("event_id", 1))
        for event in rc_events:
            if int(event.get("event_id", -1)) == int(event_id):
                return event
        return rc_events[0]
    bundle = load_result_bundle(dataset)
    events = event_intervals(bundle["label"], bundle["times"])
    if not events:
        return {"event_id": 1, "aligned_start_index": 0, "aligned_end_index": min(20, len(bundle["times"]) - 1)}
    if event_id is None:
        return events[0]
    for event in events:
        if event["event_id"] == event_id:
            return event
    return events[0]


def overview(dataset: str) -> dict[str, Any]:
    bundle = load_result_bundle(dataset)
    scores = bundle["joint_score"]
    threshold = threshold_from_summary(bundle["summary"], scores)
    times = bundle["times"]
    labels = bundle["label"][times].astype(int)
    events = root_cause_events(dataset) or event_intervals(bundle["label"], times)
    peak_idx = int(np.argmax(scores))
    metrics = bundle["summary"].get("joint_node_edge", {})
    best = metrics.get("best_f1_metrics", {})
    return {
        "dataset": dataset,
        "result_tag": bundle["result_dir"].name,
        "threshold": threshold,
        "current_time": int(times[peak_idx]),
        "current_score": _safe_float(scores[peak_idx]),
        "alert": bool(scores[peak_idx] >= threshold),
        "num_events": len(events),
        "events": events[:20],
        "metrics": {
            "point_f1": _safe_float(best.get("f1")),
            "roc_auc": _safe_float(metrics.get("roc_auc")),
            "pr_auc": _safe_float(metrics.get("pr_auc")),
            "threshold_mode": metrics.get("threshold_mode", "train_q995"),
        },
        "series": _downsample(
            [
                {
                    "time": int(t),
                    "score": _safe_float(s),
                    "node": _safe_float(n),
                    "edge": _safe_float(e),
                    "label": int(l),
                }
                for t, s, n, e, l in zip(times, scores, bundle["node_score"], bundle["edge_score"], labels)
            ]
        ),
    }


def timeseries(dataset: str, start: int | None = None, end: int | None = None) -> dict[str, Any]:
    bundle = load_result_bundle(dataset)
    test = bundle["test"]
    columns = bundle["columns"]
    times = bundle["times"]
    joint = bundle["joint_score"]
    label = bundle["label"]
    if start is None:
        start = int(max(0, times[int(np.argmax(joint))] - 700))
    if end is None:
        end = int(min(len(test) - 1, start + 1600))
    chosen = list(range(min(5, test.shape[1])))
    rc = selected_event(dataset)
    top_indices = parse_indices(rc.get("top10_joint_indices", ""))[:3]
    for idx in top_indices:
        if 0 <= idx < test.shape[1] and idx not in chosen:
            chosen.append(idx)
    raw_points = []
    for t in range(max(0, start), min(end, len(test) - 1), max(1, (end - start) // 280)):
        row = {"time": int(t), "label": int(label[t])}
        for idx in chosen[:8]:
            row[columns[idx] if idx < len(columns) else f"V{idx}"] = _safe_float(test[t, idx])
        raw_points.append(row)
    score_points = [
        {"time": int(t), "score": _safe_float(s)}
        for t, s in zip(times, joint)
        if start <= int(t) <= end
    ]
    return {
        "dataset": dataset,
        "start": start,
        "end": end,
        "sensors": [columns[i] if i < len(columns) else f"V{i}" for i in chosen[:8]],
        "points": _downsample(raw_points, 300),
        "scores": _downsample(score_points, 300),
        "events": [e for e in overview(dataset)["events"] if int(e.get("raw_end_time", e.get("end", 0))) >= start],
    }


def parse_indices(raw: Any) -> list[int]:
    if raw is None:
        return []
    if isinstance(raw, list):
        return [int(x) for x in raw]
    return [int(x) for x in str(raw).replace(",", ";").split(";") if str(x).strip().isdigit()]


def parse_names(raw: Any) -> list[str]:
    if raw is None:
        return []
    if isinstance(raw, list):
        return [str(x) for x in raw]
    return [x.strip() for x in str(raw).split(";") if x.strip()]


def root_cause(dataset: str, event_id: int | None = None) -> dict[str, Any]:
    bundle = load_result_bundle(dataset)
    columns = bundle["columns"]
    event = selected_event(dataset, event_id)
    names = parse_names(event.get("top10_joint"))
    indices = parse_indices(event.get("top10_joint_indices"))
    joint_scores = [_safe_float(x) for x in str(event.get("top10_joint_scores", "")).split(";") if x.strip()]
    node_scores = [_safe_float(x) for x in str(event.get("top10_node_scores", "")).split(";") if x.strip()]
    edge_scores = [_safe_float(x) for x in str(event.get("top10_edge_scores", "")).split(";") if x.strip()]
    if not names:
        node_errors = bundle["node_errors"]
        mean_errors = node_errors.mean(axis=0)
        indices = np.argsort(mean_errors)[::-1][:10].astype(int).tolist()
        names = [columns[i] for i in indices]
        joint_scores = [_safe_float(mean_errors[i]) for i in indices]
        node_scores = joint_scores
        edge_scores = [0.0] * len(indices)
    max_score = max(joint_scores or [1.0])
    candidates = []
    for rank, name in enumerate(names[:10], start=1):
        candidates.append(
            {
                "rank": rank,
                "name": name,
                "index": indices[rank - 1] if rank - 1 < len(indices) else None,
                "score": _safe_float(joint_scores[rank - 1] if rank - 1 < len(joint_scores) else max_score / rank),
                "normalized": _safe_float((joint_scores[rank - 1] if rank - 1 < len(joint_scores) else max_score / rank) / (max_score + 1e-9)),
                "node_score": _safe_float(node_scores[rank - 1] if rank - 1 < len(node_scores) else 0.0),
                "edge_score": _safe_float(edge_scores[rank - 1] if rank - 1 < len(edge_scores) else 0.0),
            }
        )
    top = candidates[0] if candidates else {"name": "unknown", "score": 0.0}
    return {
        "dataset": dataset,
        "event": event,
        "candidates": candidates,
        "evidence": [
            {"label": "节点预测误差", "value": f"{top['name']} error signal", "severity": "high"},
            {"label": "关系退化证据", "value": f"{top['name']} adjacent edge degradation", "severity": "high"},
            {"label": "异常窗口", "value": f"event #{event.get('event_id', 1)}", "severity": "medium"},
            {"label": "建议优先级", "value": f"优先检查 {top['name']} 及相邻变量", "severity": "medium"},
        ],
    }


def relation_graph(dataset: str, event_id: int | None = None) -> dict[str, Any]:
    rc = root_cause(dataset, event_id)
    candidates = rc["candidates"][:6]
    nodes = []
    for idx, c in enumerate(candidates):
        angle = 2 * math.pi * idx / max(len(candidates), 1)
        nodes.append(
            {
                "id": c["name"],
                "label": c["name"],
                "score": c["normalized"],
                "x": 410 + 230 * math.cos(angle),
                "y": 250 + 160 * math.sin(angle),
            }
        )
    edges = []
    for idx in range(max(0, len(nodes) - 1)):
        strength = _safe_float(candidates[idx].get("edge_score", 0.4), 0.4)
        edges.append(
            {
                "source": nodes[idx]["id"],
                "target": nodes[(idx + 1) % len(nodes)]["id"],
                "degradation": min(1.0, strength),
                "label": f"{strength:.2f}",
            }
        )
    if len(nodes) > 3:
        edges.append({"source": nodes[0]["id"], "target": nodes[3]["id"], "degradation": 0.66, "label": "0.66"})
    top_edges = sorted(edges, key=lambda e: e["degradation"], reverse=True)[:5]
    normal = [0.32, 0.44, 0.58, 0.21]
    anomaly = [0.76, 0.18, 0.71, 0.49]
    return {
        "dataset": dataset,
        "event": rc["event"],
        "nodes": nodes,
        "edges": edges,
        "top_edges": top_edges,
        "edge_vector_compare": [
            {"component": "corr", "normal": normal[0], "anomaly": anomaly[0]},
            {"component": "|corr|", "normal": normal[1], "anomaly": anomaly[1]},
            {"component": "lag", "normal": normal[2], "anomaly": anomaly[2]},
            {"component": "direction", "normal": normal[3], "anomaly": anomaly[3]},
        ],
    }


def report(dataset: str, event_id: int | None = None) -> dict[str, Any]:
    ov = overview(dataset)
    rc = root_cause(dataset, event_id)
    graph = relation_graph(dataset, event_id)
    event = rc["event"]
    top = rc["candidates"][0] if rc["candidates"] else {"name": "unknown"}
    window = f"{event.get('raw_start_time', event.get('start', '-'))}~{event.get('raw_end_time', event.get('end', '-'))}"
    return {
        "event_id": event.get("event_id", 1),
        "dataset": dataset,
        "time_window": window,
        "title": f"{dataset} 异常事件 #{event.get('event_id', 1)} 诊断报告",
        "sections": [
            {"title": "异常概况", "body": f"系统在 {window} 窗口内检测到持续异常，峰值异常分数为 {ov['current_score']:.2f}，阈值为 {ov['threshold']:.2f}。"},
            {"title": "根因候选", "body": f"{top['name']} 排名第一，候选排序由节点预测误差和相邻关系退化证据共同决定。"},
            {"title": "关系退化", "body": f"Top 退化边为 {graph['top_edges'][0]['source']} → {graph['top_edges'][0]['target']}，退化强度约 {graph['top_edges'][0]['degradation']:.2f}。"},
            {"title": "运维建议", "body": f"优先检查 {top['name']} 的读数、执行状态、相邻阀门/泵和控制链路，并复核该异常窗口的上下游联动关系。"},
        ],
    }


def agent_answer(
    dataset: str,
    question: str,
    event_id: int | None = None,
    ocr_text: str | None = None,
) -> dict[str, Any]:
    calls = ["detect_event", "rank_root_cause", "inspect_edge_degradation", "generate_report"]
    rep = report(dataset, event_id)
    rc = root_cause(dataset, event_id)
    graph = relation_graph(dataset, event_id)
    ov = overview(dataset)

    context: dict[str, Any] = {
        "overview": ov,
        "root_cause": rc,
        "relation_graph": graph,
        "report": rep,
    }
    if ocr_text:
        context["ocr_text"] = ocr_text
        calls.append("ocr_extract")

    try:
        from ernie_service import get_ernie

        ernie = get_ernie()
        result = ernie.chat(question, context)
        if result["success"]:
            calls.append("erniebot_reasoning")
            return {
                "answer": result["answer"],
                "model": result["model"],
                "tool_calls": [{"name": name, "status": "ok"} for name in calls],
                "report": rep,
            }
        logger.warning("ErnieBot 调用失败，回退到模板回答: %s", result.get("error"))
    except Exception as exc:
        logger.warning("ErnieBot 不可用，回退到模板回答: %s", exc)

    top = rc["candidates"][0] if rc["candidates"] else {"name": "unknown"}
    edge = graph["top_edges"][0] if graph["top_edges"] else {"source": top["name"], "target": "adjacent", "degradation": 0}
    answer = (
        f"该窗口报警是因为联合异常分数持续超过阈值，并且根因候选 {top['name']} 排名最高。"
        f"主要证据包括：节点预测误差升高、{edge['source']} → {edge['target']} 的关系退化强度约为 {edge['degradation']:.2f}。"
        f"建议优先检查 {top['name']} 及其相邻传感器/执行器，再核对控制链路和现场工况记录。"
    )
    if "报告" in question:
        answer = "已生成诊断报告：" + "；".join(section["body"] for section in rep["sections"])
    elif "步骤" in question or "排查" in question:
        answer = f"排查步骤：1. 核对 {top['name']} 原始读数；2. 检查 {edge['source']} 与 {edge['target']} 的联动关系；3. 复核泵/阀门状态；4. 将事件窗口导出给运维人员确认。"

    return {
        "answer": answer,
        "model": "template",
        "tool_calls": [{"name": name, "status": "ok"} for name in calls],
        "report": rep,
    }


def extract_docx_text(file_base64: str) -> dict[str, Any]:
    """从 base64 编码的 DOCX 文档提取文字（python-docx，纯 Python）。"""
    import base64
    import tempfile

    try:
        _, data = ("", file_base64)
        if "," in file_base64:
            _, data = file_base64.split(",", 1)
        file_bytes = base64.b64decode(data)
    except Exception as exc:
        return {"success": False, "text": "", "error": f"Base64 解码失败: {exc}"}

    tmp = None
    try:
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.write(file_bytes)
        tmp.close()

        from docx import Document

        doc = Document(tmp.name)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        text = "\n".join(paragraphs)
        return {"success": True, "text": text, "paragraphs_count": len(paragraphs), "error": None}
    except ImportError:
        return {"success": False, "text": "", "error": "python-docx 未安装，请执行: pip install python-docx"}
    except Exception as exc:
        logger.exception("DOCX 提取失败")
        return {"success": False, "text": "", "error": str(exc)}
    finally:
        if tmp:
            Path(tmp.name).unlink(missing_ok=True)


def extract_pdf_text(file_base64: str) -> dict[str, Any]:
    """从 base64 编码的 PDF 文档提取文字（PyPDF2）。"""
    import base64
    import tempfile

    try:
        _, data = ("", file_base64)
        if "," in file_base64:
            _, data = file_base64.split(",", 1)
        file_bytes = base64.b64decode(data)
    except Exception as exc:
        return {"success": False, "text": "", "error": f"Base64 解码失败: {exc}"}

    tmp = None
    try:
        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        tmp.write(file_bytes)
        tmp.close()

        from PyPDF2 import PdfReader

        reader = PdfReader(tmp.name)
        paragraphs = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                for line in text.split("\n"):
                    stripped = line.strip()
                    if stripped:
                        paragraphs.append(stripped)
        text = "\n".join(paragraphs)
        return {"success": True, "text": text, "paragraphs_count": len(paragraphs), "error": None}
    except ImportError:
        return {"success": False, "text": "", "error": "PyPDF2 未安装，请执行: pip install PyPDF2"}
    except Exception as exc:
        logger.exception("PDF 提取失败")
        return {"success": False, "text": "", "error": str(exc)}
    finally:
        if tmp:
            Path(tmp.name).unlink(missing_ok=True)


def ocr_extract_image(image_base64: str) -> dict[str, Any]:
    """对 base64 编码的工厂文档图片执行 OCR 文字提取（DB + SVTR_LCNet）。"""
    import base64

    try:
        _, data = ("", image_base64)
        if "," in image_base64:
            _, data = image_base64.split(",", 1)
        image_bytes = base64.b64decode(data)
    except Exception as exc:
        return {"success": False, "text": "", "items": [], "error": f"Base64 解码失败: {exc}"}

    try:
        from ocr_service import extract_text_from_bytes
        return extract_text_from_bytes(image_bytes)
    except ImportError:
        return {"success": False, "text": "", "items": [], "error": "PaddleOCR 未安装"}
    except Exception as exc:
        logger.exception("OCR 提取异常")
        return {"success": False, "text": "", "items": [], "error": str(exc)}


def extract_industrial_info(doc_text: str) -> dict[str, Any]:
    """通过 ErnieBot 从文档文本中抽取工业生产关键信息。"""
    try:
        from ernie_service import get_ernie

        return get_ernie().extract_industrial_info(doc_text)
    except ImportError:
        return {"success": False, "info": {}, "error": "ErnieBot 未安装"}
    except Exception as exc:
        logger.exception("工业信息抽取失败")
        return {"success": False, "info": {}, "error": str(exc)}


def cross_modal_analyze(
    dataset: str,
    doc_text: str,
    doc_info: dict[str, Any],
    event_id: int | None = None,
) -> dict[str, Any]:
    """跨模态关联分析：将质检文档信息与传感器异常数据关联。"""
    rc = root_cause(dataset, event_id)
    graph = relation_graph(dataset, event_id)
    ov = overview(dataset)
    rep = report(dataset, event_id)

    diagnosis_context: dict[str, Any] = {
        "overview": ov,
        "root_cause": rc,
        "relation_graph": graph,
        "report": rep,
    }

    try:
        from ernie_service import get_ernie

        return get_ernie().cross_modal_analyze(doc_text, doc_info, diagnosis_context)
    except ImportError:
        return {"success": False, "analysis": "", "error": "ErnieBot 未安装"}
    except Exception as exc:
        logger.exception("跨模态分析失败")
        return {"success": False, "analysis": "", "error": str(exc)}


# ---------- 诊断任务（/api/diagnosis/tasks） ----------

@dataclass
class DiagnosisTask:
    task_id: str
    dataset: str
    event_id: int | None
    question: str
    status: str = "created"
    stage: str = "created"
    thinking_chunks: list[str] = field(default_factory=list)
    report_chunks: list[str] = field(default_factory=list)
    tool_calls: list[dict[str, Any]] = field(default_factory=list)
    result: dict[str, Any] | None = None
    error: str | None = None
    created_at: float = field(default_factory=time.time)

DIAGNOSIS_TASKS: dict[str, DiagnosisTask] = {}
DIAGNOSIS_LOCK = threading.Lock()


def create_diagnosis_task(dataset: str, event_id: int | None, question: str) -> DiagnosisTask:
    task_id = uuid.uuid4().hex
    task = DiagnosisTask(task_id=task_id, dataset=dataset, event_id=event_id, question=question)
    with DIAGNOSIS_LOCK:
        DIAGNOSIS_TASKS[task_id] = task
    thread = threading.Thread(target=_run_diagnosis, args=(task_id,), daemon=True)
    thread.start()
    return task


def get_diagnosis_task(task_id: str) -> dict[str, Any]:
    with DIAGNOSIS_LOCK:
        task = DIAGNOSIS_TASKS.get(task_id)
    if not task:
        raise KeyError(task_id)
    return {
        "task_id": task.task_id,
        "dataset": task.dataset,
        "event_id": task.event_id,
        "question": task.question,
        "status": task.status,
        "stage": task.stage,
        "thinking_chunks": task.thinking_chunks,
        "report_chunks": task.report_chunks,
        "tool_calls": task.tool_calls,
        "result": task.result,
        "error": task.error,
    }


def get_diagnosis_chunks(task_id: str) -> dict[str, list[str]]:
    with DIAGNOSIS_LOCK:
        task = DIAGNOSIS_TASKS.get(task_id)
    if not task:
        raise KeyError(task_id)
    return {
        "thinking_chunks": task.thinking_chunks,
        "report_chunks": task.report_chunks,
    }


def _run_diagnosis(task_id: str) -> None:
    with DIAGNOSIS_LOCK:
        task = DIAGNOSIS_TASKS.get(task_id)
    if not task:
        return

    stage_messages = [
        ("detecting", "检测报警事件并读取 Relation-EVGAT 联合异常分数。"),
        ("evidence_collecting", "汇总节点预测误差、根因候选和关系退化边。"),
        ("retrieving_knowledge", "检索知识库中的变量说明、SOP 和方法资料。"),
        ("reasoning", "按 ReAct 风格串联工具结果，正在调用 ErnieBot 大模型推理。"),
        ("reporting", "生成结构化诊断报告和排查建议。"),
    ]
    try:
        for stage, message in stage_messages:
            task.status = "running"
            task.stage = stage
            task.thinking_chunks.append(message)
            time.sleep(0.08)

        result = agent_answer(task.dataset, task.question, task.event_id)
        task.result = result
        task.tool_calls = result.get("tool_calls", [])
        task.report_chunks = [result["answer"]]
        task.thinking_chunks.append("[completed] 诊断完成")
        task.stage = "completed"
        task.status = "completed"
    except Exception as exc:
        task.status = "failed"
        task.stage = "failed"
        task.error = str(exc)
        task.thinking_chunks.append(f"[error] {exc}")


# ---------- 知识库（/api/knowledge/*） ----------

def knowledge_documents() -> dict[str, Any]:
    from knowledge_service import list_documents

    return list_documents()


def knowledge_upload(filename: str, content: str) -> dict[str, Any]:
    from knowledge_service import upload_document

    return upload_document(filename, content)


def knowledge_search(query: str, top_k: int = 5) -> dict[str, Any]:
    from knowledge_service import search_knowledge

    return search_knowledge(query, top_k)


def health() -> dict[str, Any]:
    datasets = available_datasets()
    return {
        "ok": bool(datasets) and RELATION_SCRIPT.exists(),
        "project_root": str(PROJECT_ROOT),
        "relation_script": RELATION_SCRIPT.exists(),
        "datasets": datasets,
    }


def _run_training(job_id: str, dataset: str, params: dict[str, Any]) -> None:
    with JOBS_LOCK:
        job = JOBS[job_id]
        job.status = "running"
        job.updated_at = time.time()
    jobs_dir = OUTPUT_ROOT / "jobs"
    jobs_dir.mkdir(parents=True, exist_ok=True)
    log_path = jobs_dir / f"{job_id}.log"
    tag = f"_job_{job_id[:8]}"
    cmd = [
        sys.executable,
        str(RELATION_SCRIPT),
        "--project",
        str(PROJECT_ROOT),
        "--dataset",
        dataset,
        "--epochs",
        str(params.get("epochs", 1)),
        "--max-train-windows",
        str(params.get("max_train_windows", 1000)),
        "--eval-stride",
        str(params.get("eval_stride", 8)),
        "--edge-mode",
        str(params.get("edge_mode", "full")),
        "--output-tag",
        tag,
    ]
    if not params.get("use_relation_degradation", True):
        cmd.append("--no-relation-degradation")
    try:
        with log_path.open("w", encoding="utf-8") as log:
            log.write(" ".join(cmd) + "\n\n")
            proc = subprocess.run(cmd, cwd=str(PROJECT_ROOT / "relation_evgat"), text=True, stdout=log, stderr=subprocess.STDOUT)
        output_name = f"{params.get('edge_mode', 'full')}_{'joint' if params.get('use_relation_degradation', True) else 'node'}{tag}"
        with JOBS_LOCK:
            job = JOBS[job_id]
            job.status = "succeeded" if proc.returncode == 0 else "failed"
            job.log_path = str(log_path)
            job.output_dir = str(OUTPUT_ROOT / "top_ready_relation_gat" / dataset / output_name)
            job.error = None if proc.returncode == 0 else f"Training exited with code {proc.returncode}"
            job.updated_at = time.time()
    except Exception as exc:
        with JOBS_LOCK:
            job = JOBS[job_id]
            job.status = "failed"
            job.log_path = str(log_path)
            job.error = str(exc)
            job.updated_at = time.time()


def create_train_job(dataset: str, params: dict[str, Any]) -> TrainJob:
    if not (DATA_ROOT / dataset).exists():
        raise FileNotFoundError(f"Dataset not found: {dataset}")
    job_id = uuid.uuid4().hex
    job = TrainJob(job_id=job_id, dataset=dataset)
    with JOBS_LOCK:
        JOBS[job_id] = job
    thread = threading.Thread(target=_run_training, args=(job_id, dataset, params), daemon=True)
    thread.start()
    return job


def get_job(job_id: str) -> dict[str, Any]:
    with JOBS_LOCK:
        job = JOBS.get(job_id)
    if not job:
        raise KeyError(job_id)
    log_tail = ""
    if job.log_path and Path(job.log_path).exists():
        text = Path(job.log_path).read_text(encoding="utf-8", errors="ignore")
        log_tail = text[-5000:]
    return {
        "job_id": job.job_id,
        "dataset": job.dataset,
        "status": job.status,
        "created_at": job.created_at,
        "updated_at": job.updated_at,
        "log_path": job.log_path,
        "output_dir": job.output_dir,
        "error": job.error,
        "log_tail": log_tail,
    }
